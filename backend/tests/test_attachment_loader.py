# -*- coding: utf-8 -*-
"""attachment_loader.py 全量用例：附件 base64 → langchain Document → 上下文文本。

覆盖：
  - load_attachment：文本/txt 自动编码探测失败回退原始字节/图片/过大/全零/
    空数据/解析失败 error 兜底/截断 metadata
  - _dispatcher：text mime / 文本扩展名 / pdf / docx / xlsx / image / unsupported
  - _DocxLoader / _XlsxLoader 自定义 loader（真实生成 docx/xlsx 再读回）
  - attachment_context_text：预算耗尽、图片提示、过大提示、错误提示、无文本提示、
    多附件聚合
运行：pytest tests/test_attachment_loader.py
"""
import base64
import os
import sys

if __package__ in (None, ""):
    sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(__file__))))

import pytest

import app.agent.attachment_loader as al


def _b64(text: str) -> str:
    return base64.b64encode(text.encode("utf-8")).decode("ascii")


def _b64bytes(b: bytes) -> str:
    return base64.b64encode(b).decode("ascii")


# ── load_attachment ────────────────────────────────────────────────────────

def test_load_text():
    docs = al.load_attachment({"filename": "a.txt", "mime_type": "text/plain", "data": _b64("hello file")})
    assert len(docs) == 1
    assert docs[0].page_content == "hello file"
    assert docs[0].metadata["filename"] == "a.txt"
    assert docs[0].metadata["mime_type"] == "text/plain"


def test_load_text_truncated_over_6k():
    content = "x" * 7000
    docs = al.load_attachment({"filename": "big.md", "mime_type": "text/markdown", "data": _b64(content)})
    assert docs[0].metadata["truncated"] is True
    assert len(docs[0].page_content) <= al._MAX_INLINE_CHARS + 100


def test_load_image_doc():
    docs = al.load_attachment({"filename": "p.png", "mime_type": "image/png", "data": _b64bytes(b"\x89PNG")})
    assert docs[0].metadata["image"] is True


def test_load_too_large():
    # base64 解码后 > 50MB → too_large（构造大 payload 只做 size 判定，不落盘）
    docs = al.load_attachment({"filename": "huge.bin", "mime_type": "application/octet-stream",
                               "data": _b64bytes(b"\x00" * (al._MAX_TEMP_BYTES + 1))})
    assert docs[0].metadata.get("too_large") is True


def test_load_all_zero_bytes():
    docs = al.load_attachment({"filename": "zero.pdf", "mime_type": "application/pdf",
                               "data": _b64bytes(b"\x00\x00\x00\x00")})
    assert docs[0].metadata.get("empty") is True


def test_load_parse_error_fallback():
    # 非法的 docx（垃圾字节）→ error metadata，不崩溃
    docs = al.load_attachment({"filename": "broken.docx",
                               "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               "data": _b64bytes(b"\x50\x4b\x03\x04 garbage")})
    assert docs[0].metadata.get("error")
    assert docs[0].page_content == ""


def test_load_empty_data_text():
    docs = al.load_attachment({"filename": "empty.txt", "mime_type": "text/plain", "data": ""})
    assert docs[0].page_content == ""


# ── _dispatcher ────────────────────────────────────────────────────────────

def test_dispatcher_text_mime(tmp_path):
    p = tmp_path / "x.txt"
    p.write_text("hello dispatcher", encoding="utf-8")
    docs = al._dispatcher(str(p), "text/plain")
    assert docs[0].page_content == "hello dispatcher"


def test_dispatcher_binary_decode_fallback(tmp_path):
    p = tmp_path / "bad.txt"
    # 纯无效 UTF-8/GBK 字节，逼 TextLoader 自动探测失败 → 原始字节 replace 解码
    p.write_bytes(b"\xff\xfe\x01\x00\x80\x81\x82\x83")
    docs = al._dispatcher(str(p), "application/octet-stream")
    assert "�" in docs[0].page_content or docs[0].page_content != ""


def test_dispatcher_pdf(tmp_path):
    from reportlab.pdfgen import canvas
    p = tmp_path / "doc.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(72, 720, "Hello PDF")
    c.save()
    docs = al._dispatcher(str(p), "application/pdf")
    assert isinstance(docs, list)


def test_dispatcher_image_ext():
    docs = al._dispatcher("x.svg", "")
    assert docs[0].metadata["image"] is True


def test_dispatcher_unsupported():
    docs = al._dispatcher("x.xyz", "")
    assert docs[0].metadata["unsupported"] is True


def test_dispatcher_json_mime(tmp_path):
    p = tmp_path / "x.json"
    p.write_text('{"k": "v"}', encoding="utf-8")
    docs = al._dispatcher(str(p), "application/json")
    assert "v" in docs[0].page_content


# ── 自定义 loader ──────────────────────────────────────────────────────────

def test_docx_loader_paragraphs_and_tables(tmp_path):
    from docx import Document as DocxDoc
    p = tmp_path / "sample.docx"
    doc = DocxDoc()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "a"
    tbl.cell(0, 1).text = "b"
    doc.save(str(p))
    docs = al._DocxLoader(str(p)).load()
    assert "第一段" in docs[0].page_content
    assert "第二段" in docs[0].page_content
    assert "a | b" in docs[0].page_content


def test_xlsx_loader_sheets(tmp_path):
    from openpyxl import Workbook
    p = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["h1", "h2"])
    ws.append(["v1", "v2"])
    wb.save(str(p))
    docs = al._XlsxLoader(str(p)).load()
    assert "[工作表: Sheet1]" in docs[0].page_content
    assert "h1 | h2" in docs[0].page_content


# ── attachment_context_text ────────────────────────────────────────────────

def test_context_text_normal_multiple():
    files = [
        {"filename": "a.txt", "mime_type": "text/plain", "data": _b64("one")},
        {"filename": "b.txt", "mime_type": "text/plain", "data": _b64("two")},
    ]
    out = al.attachment_context_text(files, budget=6000)
    assert out == "one\n\ntwo"


def test_context_text_image_hint():
    out = al.attachment_context_text([{"filename": "i.png", "mime_type": "image/png", "data": "x"}])
    assert "[图片附件 i.png]" in out


def test_context_text_too_large_hint():
    out = al.attachment_context_text([{
        "filename": "h.bin", "mime_type": "application/octet-stream",
        "data": _b64bytes(b"\x00" * (al._MAX_TEMP_BYTES + 1)),
    }])
    assert "过大" in out


def test_context_text_error_hint():
    out = al.attachment_context_text([{
        "filename": "broken.docx",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "data": _b64bytes(b"\x50\x4b\x03\x04 garbage"),
    }])
    assert "无法解析" in out


def test_context_text_empty_hint():
    out = al.attachment_context_text([{"filename": "empty.txt", "mime_type": "text/plain", "data": ""}])
    assert "无可用文本内容" in out


def test_context_text_budget_exhausted():
    files = [
        {"filename": f"f{i}.txt", "mime_type": "text/plain", "data": _b64("y" * 3000)}
        for i in range(5)
    ]
    out = al.attachment_context_text(files, budget=3500)
    assert "[后续附件因上下文长度限制未读取]" in out


def test_context_text_trims_to_budget():
    out = al.attachment_context_text(
        [{"filename": "big.txt", "mime_type": "text/plain", "data": _b64("z" * 2000)}],
        budget=100,
    )
    assert len(out) <= 100


def test_constants():
    assert ".txt" in al._TEXT_EXTENSIONS
    assert ".png" in al._IMAGE_EXTENSIONS
    assert al._MAX_INLINE_CHARS == 6000