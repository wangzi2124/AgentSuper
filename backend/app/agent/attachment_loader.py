"""聊天附件文档读取（LangChain document loaders）。

把聊天附件（FileContent：base64 字节，内存态，无落盘文件）解析为
``langchain_core.documents.Document`` 列表，供上层内联进 LLM 上下文或入库。

策略：
- 优先用 ``langchain_community`` 官方 loader：``TextLoader``（txt/md/csv/json 等
  纯文本）、``PyPDFLoader``（pdf，需 pypdf）。
- 对官方 loader 需额外依赖（docx2txt / unstructured）但本仓库已有
  python-docx / openpyxl 的格式（docx / xlsx），用自定义 ``BaseLoader`` 包装，
  产出与 LangChain 一致的 ``Document``（page_content + metadata）。
- 未知 / 二进制 / 解析失败的格式：返回一个仅含文件名提示的 ``Document``，
  绝不让某一份附件拖垮整次对话（退化不崩溃）。

注意：langchain-community 已被官方 sunset，仅作轻量复用；如后续移除，将本模块
的自定义 loader 提升为唯一路径即可。
"""
from __future__ import annotations

import base64
import logging
import os
import tempfile
import warnings
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

# langchain-community 已 sunset，仅抑制其 DeprecationWarning（行为不受影响）
with warnings.catch_warnings():
    warnings.simplefilter("ignore", DeprecationWarning)
    from langchain_core.documents import Document  # noqa: E402
    from langchain_community.document_loaders import (  # noqa: E402
        PyPDFLoader,
        TextLoader,
    )

_MAX_INLINE_CHARS = 6000
_MAX_TEMP_BYTES = 50 * 1024 * 1024

_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json", ".log", ".py",
                    ".js", ".ts", ".html", ".xml", ".yaml", ".yml", ".ini", ".toml"}
_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".svg", ".ico"}


class _DocxLoader:
    """基于 python-docx 的 Word 文档 loader（兼容 LangChain BaseLoader 接口）。"""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        from docx import Document as DocxDoc

        doc = DocxDoc(self.file_path)
        parts: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        content = "\n".join(parts)
        return [Document(page_content=content, metadata={"source": self.file_path})]


class _XlsxLoader:
    """基于 openpyxl 的 Excel loader（兼容 LangChain BaseLoader 接口）。"""

    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[Document]:
        from openpyxl import load_workbook

        wb = load_workbook(self.file_path, data_only=True, read_only=True)
        sheet_texts: List[str] = []
        for ws in wb.worksheets:
            rows: List[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                sheet_texts.append(f"[工作表: {ws.title}]\n" + "\n".join(rows))
        content = "\n\n".join(sheet_texts)
        return [Document(page_content=content, metadata={"source": self.file_path})]


def _dispatcher(file_path: str, mime_type: str) -> List[Document]:
    """按扩展名/类型分发到对应 LangChain loader。"""
    mime = (mime_type or "").lower()
    ext = Path(file_path).suffix.lower()

    if mime.startswith("text/") or mime in ("application/json", "application/xml",
                                            "application/yaml", "application/x-yaml"):
        return TextLoader(file_path, encoding="utf-8", autodetect_encoding=True).load()

    if ext in _TEXT_EXTENSIONS:
        try:
            return TextLoader(file_path, encoding="utf-8", autodetect_encoding=True).load()
        except Exception:
            # utf-8/gbk 自动探测失败时退回原始字节解码避免抛错
            with open(file_path, "rb") as f:
                raw = f.read()
            text = raw.decode("utf-8", errors="replace")
            return [Document(page_content=text, metadata={"source": file_path})]

    if ext == ".pdf" or mime == "application/pdf":
        return PyPDFLoader(file_path).load()

    if ext in (".docx",) or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _DocxLoader(file_path).load()

    if ext in (".xlsx",) or mime == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
        return _XlsxLoader(file_path).load()

    if ext in _IMAGE_EXTENSIONS or mime.startswith("image/"):
        # 图片走多模态 image_url，不在此处文本化
        return [Document(
            page_content="",
            metadata={"source": file_path, "image": True, "mime_type": mime},
        )]

    return [Document(
        page_content="",
        metadata={"source": file_path, "unsupported": True, "mime_type": mime},
    )]


def load_attachment(f: dict) -> List[Document]:
    """把单个附件 dict 解析为 langchain Document 列表。

    Args:
        f: FileContent 结构，含 filename / mime_type / data(base64，无 data: 前缀)。
    """
    filename = f.get("filename") or "attachment"
    data_b64 = f.get("data") or ""
    mime_type = f.get("mime_type") or ""

    docs: List[Document] = []
    tmp_path: str | None = None
    try:
        raw = base64.b64decode(data_b64, validate=False)
        if len(raw) > _MAX_TEMP_BYTES:
            return [Document(page_content="", metadata={
                "source": filename, "too_large": True, "mime_type": mime_type,
            })]
        if raw and raw[0:4] in (b"\x00\x00\x00\x00",):
            # 空/全零文件按空处理
            return [Document(page_content="", metadata={
                "source": filename, "empty": True, "mime_type": mime_type,
            })]

        ext = Path(filename).suffix.lower()
        with tempfile.NamedTemporaryFile(suffix=ext or ".bin", delete=False) as tf:
            tf.write(raw)
            tmp_path = tf.name
        docs = _dispatcher(tmp_path, mime_type)
        for d in docs:
            d.metadata = {**(d.metadata or {}), "filename": filename, "mime_type": mime_type}
            if len(d.page_content) > _MAX_INLINE_CHARS:
                d.metadata["truncated"] = True
                d.page_content = d.page_content[:_MAX_INLINE_CHARS] + f"\n...[内容过长，已截断至 {_MAX_INLINE_CHARS} 字符]"
        return docs
    except Exception as e:  # noqa: BLE001 —— 附件解析失败绝不中断对话
        logger.warning("attachment parse failed (%s): %s", filename, e)
        return [Document(page_content="", metadata={
            "source": filename, "error": str(e), "mime_type": mime_type,
        })]
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def attachment_context_text(files: list[dict], budget: int = 6000) -> str:
    """把多个附件解析为一段纯文本上下文（内联进 user 消息）。

    图片/空/无法解析项生成一行提示；文档正文聚合为文本片段。
    """
    remaining = budget
    parts: List[str] = []
    for f in files:
        mime = (f.get("mime_type") or "").lower()
        if mime.startswith("image/") or Path(f.get("filename") or "").suffix.lower() in _IMAGE_EXTENSIONS:
            parts.append(f"[图片附件 {f.get('filename','file')}]")
            continue
        docs = load_attachment(f)
        text = "\n".join(d.page_content for d in docs if d.page_content).strip()
        if not text:
            md = docs[0].metadata if docs else {}
            if md.get("too_large"):
                parts.append(f"[附件 {f.get('filename','file')} 过大，未读取内容]")
            elif md.get("error"):
                parts.append(f"[附件 {f.get('filename','file')} 无法解析：{md['error']}]")
            else:
                parts.append(f"[附件 {f.get('filename','file')} 无可用文本内容]")
            continue
        if remaining <= 0:
            parts.append("[后续附件因上下文长度限制未读取]")
            break
        chunk = text[:remaining]
        remaining -= len(chunk)
        parts.append(chunk)
    return "\n\n".join(parts)
