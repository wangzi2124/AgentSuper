"""
KB Export Plugin

Searches the knowledge base and exports results as a Word document.
"""
import json
import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PLUGIN_NAME = "kb-export"
PLUGIN_VERSION = "0.1.0"
PLUGIN_DESCRIPTION = "Search the knowledge base and export results as a Word document"


def _ensure_doc(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _make_title_page(doc: Document, title: str):
    for _ in range(4):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(68, 114, 196)
    doc.add_page_break()


def tool_export_kb_to_docx(query: str, title: str = "", top_k: int = 10) -> str:
    """Search the knowledge base and create a Word document with the results.

    Use this tool when the user wants to create a Word document containing
    information retrieved from the knowledge base. This tool searches the
    KB internally and generates a well-formatted .docx file.

    Parameters:
    - query: search query to find relevant content in the knowledge base
    - title: document title (auto-generated from query if empty)
    - top_k: number of results to include (1-20)
    """
    from app.rag.plugin_bridge import get_retriever

    retriever = get_retriever()
    if not retriever:
        return "Error: retriever not available (KB not initialized)"
    if retriever.is_empty:
        return "Error: knowledge base is empty — upload documents first"

    results = retriever.invoke(query, k=top_k)
    if not results:
        return f"No results found for query: {query}"

    doc = Document()
    _ensure_doc(doc)

    doc_title = title or f"KB Search Results: {query}"
    _make_title_page(doc, doc_title)

    # Overview section
    doc.add_heading("Search Summary", level=1)
    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph(f"Results found: {len(results)}")

    # Build sections from results grouped by source document
    source_groups: dict[str, list] = {}
    for doc_data, score in results:
        meta = doc_data["metadata"]
        src = meta.get("filename") or meta.get("document_id", "unknown")
        source_groups.setdefault(src, []).append((doc_data, score))

    for i, (doc_data, score) in enumerate(results, 1):
        meta = doc_data["metadata"]
        text = doc_data["text"]
        chapter = meta.get("chapter_title", "")
        source = meta.get("filename") or meta.get("document_id", "")
        score_pct = round(score * 100, 1)

        doc.add_heading(f"Result {i}", level=2)

        info_lines = []
        if source:
            info_lines.append(f"Source: {source}")
        if chapter:
            info_lines.append(f"Chapter: {chapter}")
        info_lines.append(f"Relevance: {score_pct}%")
        info_lines.append("")
        doc.add_paragraph("\n".join(info_lines))

        doc.add_paragraph(text)

        if i < len(results):
            doc.add_paragraph("")

    # Save
    output_dir = Path(os.getcwd()) / "data" / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c for c in doc_title if c.isascii() and (c.isalnum() or c in " _-")).strip()
    if not safe_name:
        safe_name = "kb_export"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = str(output_dir / f"{safe_name}_{ts}.docx")

    doc.save(out_path)
    return f"Document created: {out_path} ({len(results)} results)"
