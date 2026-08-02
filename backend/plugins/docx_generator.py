"""
DOCX Generator Plugin

Creates Word documents (.docx) from structured content using python-docx.
"""
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PLUGIN_NAME = "docx-generator"
PLUGIN_VERSION = "0.1.0"
PLUGIN_DESCRIPTION = "Creates Word documents (.docx) from structured content"


def _set_cell_shading(cell, color_hex: str):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def _add_table(doc, headers: list, rows: list):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(h))
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, "4472C4")

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
            if ri % 2 == 1:
                _set_cell_shading(cell, "D9E2F3")

    return table


def tool_create_docx(title: str = "Document", sections: str = "[]", output_path: str = "") -> str:
    """
    Create a Word document (.docx) with structured content.

    Parameters:
    - title: document title (displayed as centered heading)
    - sections: JSON array of section objects. Each object has:
        {"type": "heading", "text": "...", "level": 1}
        {"type": "paragraph", "text": "..."}
        {"type": "table", "headers": [...], "rows": [[...], ...]}
        {"type": "bullet_list", "items": ["...", ...]}
    - output_path: optional absolute path to save (auto-generated if empty)
    """
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # Title page elements
    for _ in range(4):
        doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(68, 114, 196)

    doc.add_page_break()

    try:
        parsed = json.loads(sections)
    except json.JSONDecodeError as e:
        return f"Error: invalid sections JSON - {e}"

    for sec in parsed:
        sec_type = sec.get("type", "paragraph")

        if sec_type == "heading":
            p = doc.add_heading(sec.get("text", ""), level=sec.get("level", 1))

        elif sec_type == "paragraph":
            p = doc.add_paragraph(sec.get("text", ""))
            p.paragraph_format.space_after = Pt(6)

        elif sec_type == "table":
            headers = sec.get("headers", [])
            rows = sec.get("rows", [])
            if headers:
                _add_table(doc, headers, rows)
                doc.add_paragraph("")

        elif sec_type == "bullet_list":
            for item in sec.get("items", []):
                p = doc.add_paragraph(item, style="List Bullet")
                p.paragraph_format.space_after = Pt(2)

    base_dir = Path(__file__).resolve().parents[1]
    output_dir = base_dir / "data" / "generated"
    output_dir.mkdir(parents=True, exist_ok=True)

    if output_path:
        from app.plugins._output import resolve_output_path
        try:
            target = resolve_output_path(output_path, ".docx")
        except ValueError as e:
            return f"Error: {e}"
    else:
        from datetime import datetime
        safe_title = "".join(c for c in title if c not in '<>:"/\\|?*').strip()
        if not safe_title:
            safe_title = "document"
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        target = output_dir / f"{safe_title}_{ts}.docx"

    doc.save(str(target))
    return f"Document created successfully: {target}"
